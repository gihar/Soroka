"""Word-рендерер протокола: канонический Markdown → редактируемый .docx.

Четвёртый канал ADR-0001: тот же канонический Markdown, что и чат/PDF,
но представлен нативными стилями Word (Heading 1/2, List Number, List Bullet),
чтобы документ можно было править в Word. В отличие от PDF, эмодзи остаются —
глифы берёт шрифт Word, а не встроенный TTF.

Тесты читают сгенерированный .docx обратно python-docx'ом и проверяют
наблюдаемое: стиль абзаца, текст, жирные runs.
"""

import io

import pytest
from docx import Document

from src.services.protocol_render.docx_renderer import (
    convert_protocol_to_docx,
    convert_protocol_to_docx_async,
)


def _render(protocol_text: str) -> Document:
    """Отрендерить протокол и открыть результат для проверки."""
    return Document(io.BytesIO(convert_protocol_to_docx(protocol_text)))


def _para_num_id(paragraph):
    """numId нумерации на уровне абзаца (``None``, если задан только стилем)."""
    pPr = paragraph._p.pPr
    if pPr is None or pPr.numPr is None or pPr.numPr.numId is None:
        return None
    return pPr.numPr.numId.val


def test_section_heading_uses_word_heading_style():
    doc = _render("## Решения\n")
    paras = [p for p in doc.paragraphs if p.text]
    assert any(
        p.text == "Решения" and p.style.name == "Heading 2" for p in paras
    ), [(p.text, p.style.name) for p in paras]


def test_document_title_uses_heading_1():
    doc = _render("# Планёрка команды\n")
    paras = [p for p in doc.paragraphs if p.text]
    assert any(
        p.text == "Планёрка команды" and p.style.name == "Heading 1" for p in paras
    ), [(p.text, p.style.name) for p in paras]


def test_bullet_item_uses_list_bullet_style():
    doc = _render("- Первый пункт\n- Второй пункт\n")
    bullets = [p for p in doc.paragraphs if p.style.name == "List Bullet"]
    assert [p.text for p in bullets] == ["Первый пункт", "Второй пункт"]


def test_numbered_item_uses_list_number_style_without_literal_number():
    doc = _render("1. Первая задача\n2. Вторая задача\n")
    numbered = [p for p in doc.paragraphs if p.style.name == "List Number"]
    # Явный «1. » снят: нумерует сам Word, иначе получилось бы «1. 1. …».
    assert [p.text for p in numbered] == ["Первая задача", "Вторая задача"]


def test_inline_bold_becomes_bold_run():
    doc = _render("Решили **важное** дело\n")
    para = next(p for p in doc.paragraphs if p.text == "Решили важное дело")
    bold_texts = [r.text for r in para.runs if r.bold]
    assert bold_texts == ["важное"], [(r.text, r.bold) for r in para.runs]


def test_bold_label_line_keeps_label_bold():
    # Строка шапки «**Дата:** 20 октября» — метка жирная, значение обычное.
    doc = _render("**Дата:** 20 октября\n")
    para = next(p for p in doc.paragraphs if "Дата" in p.text)
    bold_texts = [r.text for r in para.runs if r.bold]
    assert bold_texts == ["Дата:"], [(r.text, r.bold) for r in para.runs]
    assert para.text == "Дата: 20 октября"


def test_regular_text_becomes_normal_paragraph():
    doc = _render("Обычный абзац протокола.\n")
    texts = [(p.text, p.style.name) for p in doc.paragraphs if p.text]
    assert ("Обычный абзац протокола.", "Normal") in texts, texts


def test_emoji_preserved_in_headings_and_text():
    # В отличие от PDF, эмодзи в docx остаются — глиф даёт шрифт Word.
    doc = _render("## 🎯 Цели\n- 🚀 Запуск\n")
    joined = "\n".join(p.text for p in doc.paragraphs)
    assert "🎯" in joined and "🚀" in joined, joined


def test_horizontal_rule_is_not_literal_text():
    doc = _render("Текст\n\n---\n\nЕщё текст\n")
    assert all("---" not in p.text for p in doc.paragraphs), [
        p.text for p in doc.paragraphs
    ]


def test_subheading_uses_heading_3():
    doc = _render("### Подраздел\n")
    assert any(
        p.text == "Подраздел" and p.style.name == "Heading 3"
        for p in doc.paragraphs
    ), [(p.text, p.style.name) for p in doc.paragraphs if p.text]


_SAMPLE_PROTOCOL = """# Планёрка команды продукта

**Дата:** 20 октября 2024 · 14:30
**👥 Участники:**
Иван Петров
Мария Сидорова

## ✅ Решения
1. Запускаем бету в ноябре. Обоснование: готовность фич.
2. Бюджет Q4 согласован.

## 📌 Задачи и сроки
1. Подготовить релиз-ноты — Ответственный: Иван Петров
2. Настроить мониторинг — Ответственный: Мария Сидорова

## 💡 Ключевые выводы
- Команда готова к запуску
- Риски под контролем

## 💬 Обсуждение
**Сроки релиза**
Иван Петров: предложил ноябрь.
"""


def test_full_protocol_structure():
    doc = _render(_SAMPLE_PROTOCOL)
    structure = [(p.style.name, p.text) for p in doc.paragraphs if p.text]

    # Заголовок документа и секции — нативные стили Word.
    assert ("Heading 1", "Планёрка команды продукта") in structure
    assert ("Heading 2", "✅ Решения") in structure
    # Нумерованные секции: явные номера сняты, стиль List Number.
    decisions = [t for s, t in structure if s == "List Number"]
    assert "Запускаем бету в ноябре. Обоснование: готовность фич." in decisions
    assert "Подготовить релиз-ноты — Ответственный: Иван Петров" in decisions
    # Ключевые выводы — маркированный список.
    assert ("List Bullet", "Команда готова к запуску") in structure
    # Решения и Задачи — два блока с рестартом нумерации.
    numbered = [p for p in doc.paragraphs if p.style.name == "List Number"]
    ids = [_para_num_id(p) for p in numbered]
    assert ids[0] == ids[1] and ids[2] == ids[3] and ids[0] != ids[2]


@pytest.mark.asyncio
async def test_async_wrapper_returns_openable_docx():
    data = await convert_protocol_to_docx_async("## Решения\n1. Ок\n")
    doc = Document(io.BytesIO(data))
    assert any(p.style.name == "List Number" for p in doc.paragraphs)


def test_numbered_blocks_restart_per_section():
    doc = _render(
        "## Решения\n1. Alpha\n2. Beta\n\n## Задачи\n1. Gamma\n2. Delta\n"
    )
    numbered = [p for p in doc.paragraphs if p.style.name == "List Number"]
    num_ids = [_para_num_id(p) for p in numbered]
    # Каждый абзац несёт нумерацию на своём уровне (не из общего стиля).
    assert all(n is not None for n in num_ids), num_ids
    # Два блока → два РАЗНЫХ numId, иначе Word продолжит 1,2,3,4 сквозь секции.
    assert num_ids[0] == num_ids[1]
    assert num_ids[2] == num_ids[3]
    assert num_ids[0] != num_ids[2]
