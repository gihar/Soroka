"""Доставка протокола файлом .docx (Word) и живучий фолбэк при сбое рендера."""

import io

import pytest
from docx import Document

from src.models.processing import (
    ProcessingRequest,
    ProcessingResult,
    TranscriptionResult,
)
from src.services import result_sender


def _fake_user_service(monkeypatch, mode: str) -> None:
    import src.services.user_service as user_service_module

    class FakeUserService:
        async def get_user_by_telegram_id(self, _uid):
            class User:
                protocol_output_mode = mode

            return User()

    monkeypatch.setattr(user_service_module, "UserService", FakeUserService)


def _keyboard_datas(markup) -> set:
    if markup is None:
        return set()
    return {
        btn.callback_data
        for row in markup.inline_keyboard
        for btn in row
        if btn.callback_data
    }


def _capture_documents(monkeypatch):
    """Подменить отправку документа, вернуть список пойманных (filename, bytes)."""
    captured = []

    async def fake_send_document(bot, chat_id, **kwargs):
        input_file = kwargs["document"]
        with open(input_file.path, "rb") as f:
            captured.append((input_file.filename, f.read()))
        return object()

    monkeypatch.setattr(result_sender, "safe_send_document", fake_send_document)
    return captured


@pytest.mark.asyncio
async def test_docx_mode_delivers_word_document(monkeypatch):
    captured = _capture_documents(monkeypatch)

    ok = await result_sender.send_protocol_file(
        object(), 1,
        "# Планёрка\n\n## ✅ Решения\n1. Запускаем бету\n",
        "meeting.mp3", "docx",
    )

    assert ok is True
    assert len(captured) == 1
    filename, data = captured[0]
    assert filename.endswith(".docx")
    # Доставлен настоящий .docx: открывается python-docx'ом, стили Word на месте.
    doc = Document(io.BytesIO(data))
    assert any(p.style.name == "List Number" for p in doc.paragraphs)


@pytest.mark.asyncio
async def test_primary_delivery_in_docx_mode_sends_word_file(monkeypatch):
    from unittest.mock import AsyncMock

    captured = []

    async def fake_send_message(bot, chat_id, **kwargs):
        return object()

    async def fake_send_document(bot, chat_id, **kwargs):
        input_file = kwargs["document"]
        with open(input_file.path, "rb") as f:
            captured.append((input_file.filename, f.read(), kwargs.get("reply_markup")))
        return object()

    monkeypatch.setattr(result_sender, "safe_send_message", fake_send_message)
    monkeypatch.setattr(result_sender, "safe_send_document", fake_send_document)
    _fake_user_service(monkeypatch, "docx")

    req = ProcessingRequest(file_name="meeting.mp3", llm_provider="openai", user_id=1)
    res = ProcessingResult(
        transcription_result=TranscriptionResult(transcription="текст"),
        protocol_text="# Планёрка\n\n## ✅ Решения\n1. Запускаем бету\n",
        template_used={"name": "Дейли"},
        llm_provider_used="openai",
        llm_model_used=None,
        history_id=7,
    )
    ok = await result_sender.send_result_to_user(AsyncMock(), 1, 1, req, res)

    assert ok is True
    assert len(captured) == 1
    filename, data, markup = captured[0]
    # Режим docx из настроек → первичная доставка уходит документом .docx.
    assert filename.endswith(".docx")
    doc = Document(io.BytesIO(data))
    assert any(p.style.name == "List Number" for p in doc.paragraphs)
    # Матрица кнопок: Word уже доставлен — прячем его, оставляем PDF + перегенерацию.
    assert _keyboard_datas(markup) == {"proto_pdf_7", "proto_regen_7"}


@pytest.mark.asyncio
async def test_docx_render_failure_falls_back_to_md(monkeypatch):
    captured = _capture_documents(monkeypatch)

    async def exploding_render(protocol_text):
        raise RuntimeError("docx render failed")

    monkeypatch.setattr(
        "src.services.protocol_render.docx_renderer.convert_protocol_to_docx_async",
        exploding_render,
    )

    ok = await result_sender.send_protocol_file(
        object(), 1, "# Протокол\n\n## ✅ Решения\n1. Ок\n", "meeting.mp3", "docx",
    )

    assert ok is True
    assert len(captured) == 1
    filename, data = captured[0]
    # Рендер упал — вместо тишины уходит канонический .md.
    assert filename.endswith(".md")
    assert data.decode("utf-8").startswith("# Протокол")
