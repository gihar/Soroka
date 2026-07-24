"""Критика v7: двойная маркировка списков «- 1. …» (буллет + литеральный номер).

Живой DOCX 23.07 («Статус доработки отчета по АП»): LLM недетерминированно
совмещает общее правило «маркер '- '» с полевым «НУМЕРОВАННЫЙ "N. …"» и отдаёт
«- 1. Решение…». Все каналы рендерят буллет И номер: Telegram «• 1.», DOCX —
List Bullet с выжившим номером (ветка List Number не срабатывает), PDF — «• 1.».

Фикс в два слоя:
- детерминированный нормализатор (normalize_list_markers) в финальном проходе
  хвоста completion (новые протоколы) и на входах доставки (протоколы из
  истории, перерендер кнопками PDF/Word);
- разведение правил списков в системном промпте: маркированные — «- »,
  нумерованные — «N. » БЕЗ маркера.
"""

import io

import pytest
from docx import Document

from src.prompts.prompts import (
    FIELD_SPECIFIC_RULES,
    build_generation_system_prompt,
)
from src.services import result_sender
from src.services.protocol_render import render_protocol_messages
from src.utils.text_processing import normalize_list_markers

# ---------------------------------------------------------------------------
# Нормализатор: маркер буллета перед нумерованным пунктом снимается
# ---------------------------------------------------------------------------


def test_bulleted_number_loses_bullet():
    assert normalize_list_markers("- 1. Принят формат даты") == "1. Принят формат даты"


def test_star_marker_before_number_also_dropped():
    assert normalize_list_markers("* 2. Второе решение") == "2. Второе решение"


def test_indent_preserved():
    assert normalize_list_markers("  - 3. Вложенный пункт") == "  3. Вложенный пункт"


def test_plain_bullet_untouched():
    assert normalize_list_markers("- Обычный пункт") == "- Обычный пункт"


def test_clean_numbered_item_untouched():
    assert normalize_list_markers("1. Уже чистый пункт") == "1. Уже чистый пункт"


def test_number_without_space_after_dot_untouched():
    # «1.5 часа», «05.2025» — не номера пунктов, маркер остаётся.
    assert normalize_list_markers("- 1.5 часа на созвон") == "- 1.5 часа на созвон"
    assert normalize_list_markers("- 05.2025 — срок") == "- 05.2025 — срок"


def test_horizontal_rule_untouched():
    assert normalize_list_markers("---") == "---"


def test_idempotent():
    once = normalize_list_markers("- 1. Пункт")
    assert normalize_list_markers(once) == once


def test_fenced_code_content_untouched():
    # Цитируемые блоки (логи, код) — дословный контент, маркеры в них не трогаем.
    text = (
        "## 💬 Обсуждение\n"
        "```\n"
        "- 1. первая строка лога\n"
        "- 2. вторая строка лога\n"
        "```\n"
        "- 1. Обычный пункт после фенса\n"
    )
    expected = (
        "## 💬 Обсуждение\n"
        "```\n"
        "- 1. первая строка лога\n"
        "- 2. вторая строка лога\n"
        "```\n"
        "1. Обычный пункт после фенса\n"
    )
    assert normalize_list_markers(text) == expected


def test_unclosed_fence_protects_rest_of_document():
    # Незакрытый фенс — как в telegram_html: всё до конца считается кодом.
    text = "```\n- 1. строка лога\n- 2. ещё строка"
    assert normalize_list_markers(text) == text


def test_known_tradeoff_plain_bullet_starting_with_number_is_renumbered():
    # Осознанный компромисс: буллет-пункт, чей текст начинается с «N. »
    # («- 1. этап миграции…»), синтаксически неотличим от нумерованного
    # пункта со случайным маркером — нормализатор снимет маркер и здесь.
    assert (
        normalize_list_markers("- 1. этап миграции завершить — Отв.: Иван")
        == "1. этап миграции завершить — Отв.: Иван"
    )


def test_multiline_document_normalized_linewise():
    text = (
        "## ✅ Решения\n"
        "- 1. Принят формат даты\n"
        "- 2. Быстрые правки\n"
        "\n"
        "## 📌 Задачи и сроки\n"
        "1. Привести даты к формату\n"
        "- обычный буллет\n"
    )
    expected = (
        "## ✅ Решения\n"
        "1. Принят формат даты\n"
        "2. Быстрые правки\n"
        "\n"
        "## 📌 Задачи и сроки\n"
        "1. Привести даты к формату\n"
        "- обычный буллет\n"
    )
    assert normalize_list_markers(text) == expected


# ---------------------------------------------------------------------------
# Шов completion: новый протокол выходит из хвоста уже нормализованным
# ---------------------------------------------------------------------------


@pytest.mark.asyncio
async def test_completion_tail_normalizes_list_markers():
    from types import SimpleNamespace
    from unittest.mock import AsyncMock

    from src.models.processing import ProcessingRequest, TranscriptionResult
    from src.services.processing.completion import CompletionDeps, complete_processing

    deps = CompletionDeps(
        llm_gen=SimpleNamespace(
            optimized_llm_generation=AsyncMock(return_value={"meeting_title": "Планёрка"}),
            resolve_model_display_name=AsyncMock(return_value="GPT"),
        ),
        formatter=SimpleNamespace(
            format_protocol=lambda *a, **k: "## ✅ Решения\n- 1. Принят формат даты\n"
        ),
        history=SimpleNamespace(
            save_processing_history=AsyncMock(return_value=42),
            cleanup_temp_file=AsyncMock(),
        ),
    )

    async def delivery(result):
        return True

    outcome = await complete_processing(
        request=ProcessingRequest(file_name="a.mp3", llm_provider="openai", user_id=1),
        transcription_result=TranscriptionResult(transcription="текст"),
        template=SimpleNamespace(name="Дейли"),
        meeting_type=None,
        deps=deps,
        delivery=delivery,
    )

    assert "- 1." not in outcome.result.protocol_text
    assert "1. Принят формат даты" in outcome.result.protocol_text


# ---------------------------------------------------------------------------
# Входы доставки: протоколы из истории (уже сохранённые с дефектом) чинятся
# при перерендере — чат, .md и .docx
# ---------------------------------------------------------------------------

_LEGACY_PROTOCOL = (
    "# Статус доработки отчёта\n\n"
    "## ✅ Решения\n"
    "- 1. Принят формат даты ММ.ГГГГ\n"
    "- 2. Быстрые правки выполнить оперативно\n"
)


def test_chat_channel_renders_number_without_bullet():
    parts = render_protocol_messages(_LEGACY_PROTOCOL)
    joined = "\n".join(parts)
    assert "• 1." not in joined
    assert "1. Принят формат даты ММ.ГГГГ" in joined


def _capture_documents(monkeypatch):
    captured = []

    async def fake_send_document(bot, chat_id, **kwargs):
        input_file = kwargs["document"]
        with open(input_file.path, "rb") as f:
            captured.append((input_file.filename, f.read()))
        return object()

    monkeypatch.setattr(result_sender, "safe_send_document", fake_send_document)
    return captured


@pytest.mark.asyncio
async def test_docx_channel_uses_word_numbering_without_literal_number(monkeypatch):
    captured = _capture_documents(monkeypatch)

    ok = await result_sender.send_protocol_file(
        object(), 1, _LEGACY_PROTOCOL, "meeting.mp3", "docx",
    )

    assert ok is True
    _, data = captured[0]
    doc = Document(io.BytesIO(data))
    numbered = [p.text for p in doc.paragraphs if p.style.name == "List Number"]
    bullets = [p.text for p in doc.paragraphs if p.style.name == "List Bullet"]
    # Нумерует сам Word: литеральный «1. » снят, буллетов в решениях нет.
    assert numbered == [
        "Принят формат даты ММ.ГГГГ",
        "Быстрые правки выполнить оперативно",
    ]
    assert bullets == []


@pytest.mark.asyncio
async def test_md_channel_file_content_normalized(monkeypatch):
    captured = _capture_documents(monkeypatch)

    ok = await result_sender.send_protocol_file(
        object(), 1, _LEGACY_PROTOCOL, "meeting.mp3", "file",
    )

    assert ok is True
    filename, data = captured[0]
    assert filename.endswith(".md")
    content = data.decode("utf-8")
    assert "- 1." not in content
    assert "1. Принят формат даты ММ.ГГГГ" in content


# ---------------------------------------------------------------------------
# Промпт: правила списков разведены, противоречие «- » vs «N. » снято
# ---------------------------------------------------------------------------


def test_system_prompt_splits_bulleted_and_numbered_list_rules():
    base = build_generation_system_prompt()
    # Старое универсальное правило диктовало маркер «- » ЛЮБОМУ списку.
    assert "Списки: через" not in base
    # Нумерованные списки — явно БЕЗ маркера перед номером.
    assert "БЕЗ маркера '- '" in base
    assert "Маркированные списки" in base


def test_decisions_rule_demands_short_lead_and_human_attribution():
    rule = FIELD_SPECIFIC_RULES["decisions"]
    # Решение — первой короткой фразой, хвост не длиннее решения.
    assert "короткой фразой" in rule
    # Атрибуция человеческим оборотом, не телеграфом «Имя: глагол; Имя: глагол».
    assert "предложила Елена" in rule
    assert "Имя: глагол" in rule


def test_meeting_title_rule_bans_bare_abbreviations():
    rule = FIELD_SPECIFIC_RULES["meeting_title"]
    assert "аббревиатур" in rule


def test_base_prompt_requires_consistent_yo():
    base = build_generation_system_prompt()
    assert "«ё»" in base
