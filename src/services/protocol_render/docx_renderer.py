"""Word (.docx) рендерер протокола — четвёртый канал ADR-0001.

Канонический формат протокола — Markdown. Этот модуль представляет его нативными
стилями Word (Heading 1/2, List Number, List Bullet), чтобы документ открывался
и правился в Word как обычный текст, а не как «картинка». В отличие от PDF,
эмодзи НЕ снимаются: их глифы даёт шрифт Word.

Нумерованные списки каждой секции нумеруются самим Word (стиль List Number),
поэтому явный «1. » из Markdown снимается. Чтобы Word не продолжал нумерацию
сквозь секции (1,2,3 в «Решениях», 4,5 в «Задачах»), каждый блок получает свой
экземпляр нумерации с рестартом на 1.
"""

import io
import re

from docx import Document
from docx.oxml import OxmlElement
from docx.oxml.ns import qn

# Нумерованный пункт канонического Markdown: «1. », «2. » … Явный номер снимается.
_NUMBERED_RE = re.compile(r"^\d+\.\s+(.*)$")
# Инлайновый **жирный** — согласовано с чат-рендером и PDF: только двойные звёзды.
_BOLD_RE = re.compile(r"\*\*(.+?)\*\*")


def _add_inline_runs(paragraph, text: str) -> None:
    """Разложить `**жирный**` на runs; остальной текст — обычными runs.

    Эмодзи не трогаем (в отличие от PDF): их отображает шрифт Word.
    """
    pos = 0
    for match in _BOLD_RE.finditer(text):
        if match.start() > pos:
            paragraph.add_run(text[pos:match.start()])
        paragraph.add_run(match.group(1)).bold = True
        pos = match.end()
    if pos < len(text):
        paragraph.add_run(text[pos:])


def _list_number_abstract_id(document) -> str:
    """abstractNumId встроенного стиля List Number — переиспользуем его формат."""
    numbering = document.part.numbering_part.numbering_definitions._numbering
    style_pPr = document.styles["List Number"].element.pPr
    style_num_id = style_pPr.numPr.numId.val
    for num in numbering.findall(qn("w:num")):
        if int(num.get(qn("w:numId"))) == style_num_id:
            return num.find(qn("w:abstractNumId")).get(qn("w:val"))
    # Фолбэк: любой abstractNum, если стиль неожиданно без нумерации.
    first = numbering.find(qn("w:abstractNum"))
    return first.get(qn("w:abstractNumId")) if first is not None else "0"


def _new_restarting_num(document, abstract_num_id: str) -> int:
    """Создать новый экземпляр нумерации со стартом с 1, вернуть его numId."""
    numbering = document.part.numbering_part.numbering_definitions._numbering
    used = [int(n.get(qn("w:numId"))) for n in numbering.findall(qn("w:num"))]
    new_id = max(used, default=0) + 1

    num = OxmlElement("w:num")
    num.set(qn("w:numId"), str(new_id))
    abstract_ref = OxmlElement("w:abstractNumId")
    abstract_ref.set(qn("w:val"), str(abstract_num_id))
    num.append(abstract_ref)
    lvl_override = OxmlElement("w:lvlOverride")
    lvl_override.set(qn("w:ilvl"), "0")
    start_override = OxmlElement("w:startOverride")
    start_override.set(qn("w:val"), "1")
    lvl_override.append(start_override)
    num.append(lvl_override)
    numbering.append(num)
    return new_id


def _apply_num(paragraph, num_id: int) -> None:
    """Привязать абзац к конкретному экземпляру нумерации (ilvl 0)."""
    numPr = paragraph._p.get_or_add_pPr().get_or_add_numPr()
    numPr.get_or_add_ilvl().val = 0
    numPr.get_or_add_numId().val = num_id


def convert_protocol_to_docx(protocol_text: str) -> bytes:
    """Собрать .docx из канонического Markdown протокола, вернуть его байты."""
    document = Document()
    abstract_num_id = _list_number_abstract_id(document)
    current_num_id: int | None = None  # нумерация текущего блока; None — блока нет

    for line in protocol_text.split("\n"):
        stripped = line.strip()
        if not stripped:
            continue  # пустая строка — разделитель, но не разрыв нумерованного блока

        numbered = _NUMBERED_RE.match(stripped)
        if numbered:
            if current_num_id is None:
                current_num_id = _new_restarting_num(document, abstract_num_id)
            paragraph = document.add_paragraph(style="List Number")
            _add_inline_runs(paragraph, numbered.group(1).strip())
            _apply_num(paragraph, current_num_id)
            continue

        current_num_id = None  # любой не-нумерованный контент закрывает блок
        # Линейка-разделитель Markdown (---): в Word секции разводят стили
        # заголовков, литеральное «---» было бы мусором.
        if re.match(r"^-{3,}$", stripped):
            continue
        if stripped.startswith("#### "):
            _add_inline_runs(
                document.add_paragraph(style="Heading 4"), stripped[5:].strip()
            )
        elif stripped.startswith("### "):
            _add_inline_runs(
                document.add_paragraph(style="Heading 3"), stripped[4:].strip()
            )
        elif stripped.startswith("## "):
            _add_inline_runs(
                document.add_paragraph(style="Heading 2"), stripped[3:].strip()
            )
        elif stripped.startswith("# "):
            _add_inline_runs(
                document.add_paragraph(style="Heading 1"), stripped[2:].strip()
            )
        elif stripped.startswith("- ") or stripped.startswith("* "):
            _add_inline_runs(
                document.add_paragraph(style="List Bullet"), stripped[2:].strip()
            )
        else:
            _add_inline_runs(document.add_paragraph(), stripped)

    buffer = io.BytesIO()
    document.save(buffer)
    return buffer.getvalue()


async def convert_protocol_to_docx_async(protocol_text: str) -> bytes:
    """Async-обёртка: рендер .docx в пуле потоков, чтобы не блокировать loop."""
    from src.performance.async_optimization import thread_manager

    return await thread_manager.run_in_thread(convert_protocol_to_docx, protocol_text)
